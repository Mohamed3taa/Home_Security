import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def parse_inline(run_text):
    """Returns list of (text, bold, italic, code) tuples"""
    parts = []
    pattern = re.compile(r'(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*|[^`*]+)')
    for m in pattern.finditer(run_text):
        t = m.group(0)
        if t.startswith('`') and t.endswith('`'):
            parts.append((t[1:-1], False, False, True))
        elif t.startswith('**') and t.endswith('**'):
            parts.append((t[2:-2], True, False, False))
        elif t.startswith('*') and t.endswith('*'):
            parts.append((t[1:-1], False, True, False))
        else:
            parts.append((t, False, False, False))
    return parts

def add_paragraph_with_inline(doc, text, style=None, alignment=None):
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    parts = parse_inline(text)
    for (t, bold, italic, code) in parts:
        run = p.add_run(t)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
    return p

def convert_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            in_table = False
            return
        # Filter separator rows
        data_rows = [r for r in table_rows if not re.match(r'^\s*\|[-| :]+\|\s*$', r)]
        if not data_rows:
            in_table = False
            table_rows = []
            return

        parsed = []
        for row in data_rows:
            cells = [c.strip() for c in row.strip().strip('|').split('|')]
            parsed.append(cells)

        max_cols = max(len(r) for r in parsed)
        tbl = doc.add_table(rows=len(parsed), cols=max_cols)
        tbl.style = 'Table Grid'
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

        for ri, row in enumerate(parsed):
            for ci, cell_text in enumerate(row):
                if ci >= max_cols:
                    break
                cell = tbl.cell(ri, ci)
                cell.text = ''
                p = cell.paragraphs[0]
                parts = parse_inline(cell_text)
                for (t, bold, italic, code) in parts:
                    run = p.add_run(t)
                    run.bold = bold or (ri == 0)
                    run.italic = italic
                    run.font.size = Pt(10)
                    if code:
                        run.font.name = 'Courier New'
                if ri == 0:
                    set_cell_bg(cell, 'D6E4F0')

        doc.add_paragraph()
        in_table = False
        table_rows = []

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Code block
        if line.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines = []
                i += 1
                continue
            else:
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1)
                    for cl in code_lines:
                        run = p.add_run(cl + '\n')
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
                    p.paragraph_format.space_after = Pt(6)
                i += 1
                continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table detection
        if line.strip().startswith('|'):
            if in_table:
                table_rows.append(line)
            else:
                in_table = True
                table_rows = [line]
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        stripped = line.strip()

        # Horizontal rule
        if stripped in ('---', '***', '___'):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Empty line
        if stripped == '':
            i += 1
            continue

        # Headings
        if stripped.startswith('#### '):
            p = add_paragraph_with_inline(doc, stripped[5:], style='Heading 4')
        elif stripped.startswith('### '):
            p = add_paragraph_with_inline(doc, stripped[4:], style='Heading 3')
        elif stripped.startswith('## '):
            p = add_paragraph_with_inline(doc, stripped[3:], style='Heading 2')
        elif stripped.startswith('# '):
            p = add_paragraph_with_inline(doc, stripped[2:], style='Heading 1')

        # Bullet list
        elif stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = add_paragraph_with_inline(doc, text, style='List Bullet')

        # Numbered list
        elif re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s', '', stripped)
            p = add_paragraph_with_inline(doc, text, style='List Number')

        # Normal paragraph
        else:
            p = add_paragraph_with_inline(doc, stripped)

        i += 1

    if in_table:
        flush_table()

    doc.save(docx_path)
    print(f"Done! Saved to: {docx_path}")

if __name__ == '__main__':
    convert_md_to_docx(
        r'd:\project\Documentation\Graduation_Project_Documentation.md',
        r'd:\project\Documentation\Graduation_Project_Documentation.docx'
    )
